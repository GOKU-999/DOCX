import streamlit as st
import os
import tempfile
from pathlib import Path
import subprocess
import sys

# Page configuration
st.set_page_config(
    page_title="DOCX to PDF Converter - No API Required",
    page_icon="📄",
    layout="centered",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 1rem;
    }
    .success-msg {
        padding: 1rem;
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        border-radius: 0.5rem;
        color: #155724;
    }
    .error-msg {
        padding: 1rem;
        background-color: #f8d7da;
        border: 1px solid #f5c6cb;
        border-radius: 0.5rem;
        color: #721c24;
    }
    .upload-section {
        border: 2px dashed #1f77b4;
        border-radius: 10px;
        padding: 2rem;
        text-align: center;
        margin: 1rem 0;
    }
    .warning-box {
        background-color: #fff3cd;
        border: 1px solid #ffeaa7;
        border-radius: 0.5rem;
        padding: 1rem;
        margin: 1rem 0;
        color: #000000 !important;
    }
    .feature-box {
        background-color: #f8f9fa;
        border: 1px solid #e9ecef;
        border-radius: 0.5rem;
        padding: 1rem;
        margin: 0.5rem 0;
    }
    .stButton button {
        width: 100%;
    }
</style>
""", unsafe_allow_html=True)

def check_libreoffice():
    """Check if LibreOffice is available on the system"""
    try:
        # Try different possible command names
        for cmd in ['libreoffice', 'soffice', 'libreoffice7.6', 'libreoffice7.5']:
            try:
                result = subprocess.run([cmd, '--version'], 
                                      capture_output=True, text=True, timeout=10)
                if result.returncode == 0:
                    return True, cmd
            except (FileNotFoundError, subprocess.TimeoutExpired):
                continue
        return False, None
    except Exception:
        return False, None

def convert_with_libreoffice(docx_file, output_path):
    """
    Convert DOCX to PDF using LibreOffice - Perfect formatting preservation
    """
    try:
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_docx:
            temp_docx.write(docx_file.getvalue())
            temp_docx_path = temp_docx.name
        
        # Get LibreOffice command
        libreoffice_available, cmd_name = check_libreoffice()
        
        if not libreoffice_available:
            return False, "LibreOffice not found. Please install LibreOffice for perfect conversion."
        
        # Convert using LibreOffice
        cmd = [
            cmd_name, '--headless', '--convert-to', 'pdf',
            '--outdir', os.path.dirname(output_path),
            temp_docx_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        
        # Clean up temporary DOCX file
        os.unlink(temp_docx_path)
        
        if result.returncode == 0:
            # LibreOffice creates PDF with same name but .pdf extension
            expected_pdf_path = temp_docx_path.replace('.docx', '.pdf')
            if os.path.exists(expected_pdf_path):
                os.rename(expected_pdf_path, output_path)
                return True, "✅ Perfect conversion using LibreOffice! All formatting preserved."
            else:
                return False, "❌ PDF file was not created by LibreOffice"
        else:
            error_msg = result.stderr if result.stderr else "Unknown error occurred"
            return False, f"❌ LibreOffice conversion failed: {error_msg}"
            
    except subprocess.TimeoutExpired:
        return False, "❌ Conversion timeout - file might be too large or complex"
    except Exception as e:
        return False, f"❌ Conversion error: {str(e)}"

def convert_with_python(docx_file, output_path):
    """
    Fallback conversion using Python libraries (text and basic formatting only)
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        import docx
        
        # Read DOCX content
        doc = docx.Document(docx_file)
        
        # Create PDF document
        doc_pdf = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Create styles
        styles = getSampleStyleSheet()
        story = []
        
        # Process each paragraph
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Detect headings and apply styles
                if paragraph.style.name.startswith('Heading'):
                    if 'Heading 1' in paragraph.style.name:
                        style = styles['Heading1']
                    else:
                        style = styles['Heading2']
                else:
                    style = styles['Normal']
                
                # Create paragraph for PDF
                p = Paragraph(paragraph.text.replace('\n', '<br/>'), style)
                story.append(p)
                story.append(Spacer(1, 6))
        
        # Build PDF
        if story:
            doc_pdf.build(story)
            return True, "✅ Basic conversion completed. Text content preserved."
        else:
            return False, "❌ No text content found in the document."
            
    except ImportError as e:
        return False, f"❌ Required Python package missing: {str(e)}"
    except Exception as e:
        return False, f"❌ Python conversion error: {str(e)}"

def get_document_preview(docx_file):
    """Extract text content for preview"""
    try:
        import docx
        doc = docx.Document(docx_file)
        content = []
        
        for i, paragraph in enumerate(doc.paragraphs[:8]):  # First 8 paragraphs
            if paragraph.text.strip():
                # Truncate very long paragraphs
                text = paragraph.text
                if len(text) > 200:
                    text = text[:200] + "..."
                content.append(f"**Paragraph {i+1}:** {text}")
                if len(content) >= 5:  # Show max 5 paragraphs
                    break
        
        return content if content else ["No readable text content found"]
    except Exception as e:
        return [f"Could not preview document: {str(e)}"]

def main():
    # Header
    st.markdown('<h1 class="main-header">📄 DOCX to PDF Converter</h1>', unsafe_allow_html=True)
    
    st.markdown("""
    <div style='text-align: center; margin-bottom: 2rem;'>
    <h3>Free • No API Keys • No Registration • Privacy First</h3>
    Convert your Microsoft Word documents to PDF format instantly using open-source technology.
    </div>
    """, unsafe_allow_html=True)
    
    # Check system capabilities
    libreoffice_available, cmd_name = check_libreoffice()
    
    if libreoffice_available:
        st.markdown("""
        <div class="success-msg">
        ✅ <strong>LibreOffice Detected!</strong> Your documents will be converted with perfect formatting preservation.
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
        <div class="warning-box">
        ⚠️ <strong>LibreOffice Not Found</strong> - Using basic text conversion. 
        For perfect formatting (tables, images, layouts), please install LibreOffice.
        </div>
        """, unsafe_allow_html=True)
    
    # File upload section
    st.markdown('<div class="upload-section">', unsafe_allow_html=True)
    uploaded_file = st.file_uploader(
        "Choose a DOCX file or drag and drop here", 
        type=['docx'],
        help="Upload any .docx file up to 200MB"
    )
    st.markdown('</div>', unsafe_allow_html=True)
    
    if uploaded_file is not None:
        # Display file info
        col1, col2, col3 = st.columns(3)
        with col1:
            st.info(f"**File:** {uploaded_file.name}")
        with col2:
            st.info(f"**Size:** {uploaded_file.size / 1024:.1f} KB")
        with col3:
            st.info(f"**Type:** {uploaded_file.type}")
        
        # Document preview
        with st.expander("📋 Document Preview (First 5 paragraphs)", expanded=True):
            preview_content = get_document_preview(uploaded_file)
            for item in preview_content:
                st.write(item)
        
        # Conversion options
        st.subheader("Conversion Options")
        
        if libreoffice_available:
            conversion_method = st.radio(
                "Select conversion quality:",
                ["High Quality (LibreOffice)", "Basic (Python)"],
                help="High Quality preserves all formatting, Basic is text-only"
            )
        else:
            conversion_method = "Basic (Python)"
            st.warning("Only basic conversion available. Install LibreOffice for high quality.")
        
        # Convert button
        if st.button("🚀 Convert to PDF", type="primary", use_container_width=True):
            with st.spinner("Converting your document... Please wait."):
                # Create temporary file for PDF output
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_pdf:
                    temp_pdf_path = temp_pdf.name
                
                # Perform conversion based on selected method
                if conversion_method == "High Quality (LibreOffice)":
                    success, message = convert_with_libreoffice(uploaded_file, temp_pdf_path)
                else:
                    success, message = convert_with_python(uploaded_file, temp_pdf_path)
                
                # Fallback if high quality fails
                if not success and conversion_method == "High Quality (LibreOffice)":
                    st.warning("High quality conversion failed, trying basic method...")
                    success, message = convert_with_python(uploaded_file, temp_pdf_path)
                
                if success:
                    st.markdown(f'<div class="success-msg">{message}</div>', unsafe_allow_html=True)
                    
                    # Read the converted PDF
                    with open(temp_pdf_path, "rb") as f:
                        pdf_data = f.read()
                    
                    # Download section
                    st.subheader("📥 Download Your PDF")
                    
                    col1, col2 = st.columns([2, 1])
                    with col1:
                        download_filename = Path(uploaded_file.name).stem + ".pdf"
                        st.download_button(
                            label="💾 Download PDF File",
                            data=pdf_data,
                            file_name=download_filename,
                            mime="application/pdf",
                            use_container_width=True
                        )
                    
                    with col2:
                        # Show file info
                        pdf_size = len(pdf_data) / 1024
                        st.metric("PDF Size", f"{pdf_size:.1f} KB")
                    
                    # Clean up temporary file
                    os.unlink(temp_pdf_path)
                    
                    st.balloons()
                    
                else:
                    st.markdown(f'<div class="error-msg">{message}</div>', unsafe_allow_html=True)

    # Features and Instructions
    col1, col2 = st.columns(2)
    
    with col1:
        with st.expander("✨ Features", expanded=True):
            st.markdown("""
            <div class="feature-box">
            🔒 <strong>No API Keys</strong><br>
            Completely free, no registration required
            </div>
            
            <div class="feature-box">
            🛡️ <strong>Privacy First</strong><br>
            Files processed locally, never uploaded to cloud
            </div>
            
            <div class="feature-box">
            🎯 <strong>Perfect Formatting</strong><br>
            With LibreOffice, preserves all formatting
            </div>
            
            <div class="feature-box">
            ⚡ <strong>Fast Conversion</strong><br>
            Quick processing even for large files
            </div>
            """, unsafe_allow_html=True)
    
    with col2:
        with st.expander("📖 How to Use", expanded=True):
            st.markdown("""
            1. **Upload** your .docx file
            2. **Preview** the document content
            3. **Select** conversion quality
            4. **Click** Convert to PDF
            5. **Download** your perfect PDF
            
            **Supported:** .docx files only
            **Max Size:** 200MB
            **Platforms:** Windows, Mac, Linux
            """)
    
    # Installation guide
    with st.expander("🔧 Install LibreOffice for Perfect Conversion"):
        st.markdown("""
        ### Windows:
        1. Download from [libreoffice.org](https://www.libreoffice.org/download/download-libreoffice/)
        2. Run the installer
        3. Restart this app
        
        ### Mac:
        ```bash
        brew install --cask libreoffice
        ```
        or download from website
        
        ### Linux:
        ```bash
        # Ubuntu/Debian
        sudo apt update && sudo apt install libreoffice
        
        # CentOS/RHEL
        sudo yum install libreoffice
        
        # Fedora
        sudo dnf install libreoffice
        ```
        
        After installation, restart this app for LibreOffice detection.
        """)
    
    # Footer
    st.markdown("---")
    st.markdown(
        "<div style='text-align: center; color: #666;'>"
        "Made with ❤️ using Open Source • No API Keys • Privacy Focused"
        "</div>",
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()
